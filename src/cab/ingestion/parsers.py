# SPDX-License-Identifier: MIT
"""Document parsers — text + lightweight metadata extraction.

Supported formats: PDF, DOCX, MD, TXT. TXT/MD parse with the standard library
(no dependency). PDF/DOCX use OPTIONAL, lazy-imported libraries (pypdf,
python-docx) so the engine + test suite run hermetically without them; if the
library is absent the parser raises ParserUnavailable, which the ingestion path
surfaces as a clear error rather than a crash.

Metadata extraction follows Dublin Core / schema.org referents (title, creator/
author, date, source) — public vocabularies, no production implementation.
"""

from __future__ import annotations

import re
from typing import Any

from cab.models import Document

SUPPORTED_EXTS = ("pdf", "docx", "md", "markdown", "txt")

# Conservative date patterns for staleness/provenance detection.
_DATE_RE = re.compile(r"\b(20\d{2})[-/](0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])\b")
_FRONTMATTER_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.DOTALL)


class ParserUnavailable(RuntimeError):
    """A format requires an optional dependency that is not installed."""


def ext_of(name: str) -> str:
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


def parse(name: str, data: bytes) -> Document:
    """Parse raw bytes into a normalized Document by file extension."""
    ext = ext_of(name)
    if ext in ("md", "markdown"):
        text = data.decode("utf-8", errors="replace")
        return _finish(name, ext, text, data, _md_metadata(text))
    if ext == "txt":
        text = data.decode("utf-8", errors="replace")
        return _finish(name, ext, text, data, _text_metadata(text))
    if ext == "pdf":
        return _finish(name, ext, _parse_pdf(data), data, {})
    if ext == "docx":
        return _finish(name, ext, _parse_docx(data), data, {})
    raise ParserUnavailable(f"unsupported format: .{ext}")


def _finish(name: str, ext: str, text: str, data: bytes, meta: dict[str, Any]) -> Document:
    meta = dict(meta)
    # Infer title from the first markdown/heading line if not declared.
    if "title" not in meta:
        m = re.search(r"^#{1,3}\s+(.+)$", text, re.MULTILINE)
        if m:
            meta["title"] = m.group(1).strip()
    # Infer a date if present anywhere (provenance/freshness signal).
    if "date" not in meta:
        d = _DATE_RE.search(text)
        if d:
            meta["date"] = d.group(0).replace("/", "-")
    return Document(
        name=name,
        ext=ext,
        text=text,
        raw_bytes_len=len(data),
        metadata=meta,
        parsed_ok=bool(text.strip()),
    )


def _md_metadata(text: str) -> dict[str, Any]:
    """Extract YAML-ish front matter (key: value) without a YAML dependency."""
    meta: dict[str, Any] = {}
    m = _FRONTMATTER_RE.match(text)
    if not m:
        return _text_metadata(text)
    for line in m.group(1).splitlines():
        if ":" in line:
            k, _, v = line.partition(":")
            k = k.strip().lower()
            v = v.strip().strip("'\"")
            if k in ("title", "author", "creator", "date", "source", "version"):
                meta[_normalize_key(k)] = v
    return meta


def _text_metadata(text: str) -> dict[str, Any]:
    """Best-effort inline metadata (Author:/Date:/Source: lines)."""
    meta: dict[str, Any] = {}
    for key in ("author", "creator", "date", "source", "version"):
        m = re.search(rf"^{key}\s*[:\-]\s*(.+)$", text, re.IGNORECASE | re.MULTILINE)
        if m:
            meta[_normalize_key(key)] = m.group(1).strip()
    return meta


def _normalize_key(k: str) -> str:
    return "author" if k == "creator" else k


def _parse_pdf(data: bytes) -> str:
    try:
        import pypdf
    except Exception as exc:  # pragma: no cover - exercised only with the dep absent
        raise ParserUnavailable("PDF parsing requires the optional 'pypdf' package") from exc
    import io  # pragma: no cover

    reader = pypdf.PdfReader(io.BytesIO(data))  # pragma: no cover
    return "\n".join((page.extract_text() or "") for page in reader.pages)  # pragma: no cover


def _parse_docx(data: bytes) -> str:
    try:
        import docx
    except Exception as exc:  # pragma: no cover - exercised only with the dep absent
        raise ParserUnavailable("DOCX parsing requires the optional 'python-docx' package") from exc
    import io  # pragma: no cover

    document = docx.Document(io.BytesIO(data))  # pragma: no cover
    return "\n".join(p.text for p in document.paragraphs)  # pragma: no cover
